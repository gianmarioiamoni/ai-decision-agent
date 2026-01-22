# app/report/docx_converter.py
#
# DOCX Converter - Converts HTML reports to DOCX format.
#

from typing import Optional
import tempfile
from html.parser import HTMLParser


class HTMLToDocxConverter(HTMLParser):
    #
    # Simple HTML to plain text converter for DOCX generation.
    #
    
    def __init__(self):
        super().__init__()
        self.text_parts = []
        self.in_style = False
        self.in_script = False
    
    def handle_starttag(self, tag, attrs):
        if tag == 'style':
            self.in_style = True
        elif tag == 'script':
            self.in_script = True
        elif tag == 'br':
            self.text_parts.append('\n')
        elif tag in ['p', 'div']:
            self.text_parts.append('\n')
    
    def handle_endtag(self, tag):
        if tag == 'style':
            self.in_style = False
        elif tag == 'script':
            self.in_script = False
        elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.text_parts.append('\n\n')
        elif tag in ['p', 'div', 'li']:
            self.text_parts.append('\n')
    
    def handle_data(self, data):
        if not self.in_style and not self.in_script:
            self.text_parts.append(data)
    
    def get_text(self):
        return ''.join(self.text_parts)


def html_to_docx(html_content: str) -> Optional[bytes]:
    #
    # Convert HTML content to DOCX bytes.
    #
    # Args:
    #     html_content: HTML string to convert
    #
    # Returns:
    #     DOCX bytes, or None if conversion fails
    #
    # Raises:
    #     ImportError: If python-docx is not installed
    #
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import io
        
        # Parse HTML to extract text
        parser = HTMLToDocxConverter()
        parser.feed(html_content)
        text_content = parser.get_text()
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading('AI Decision Session Report', level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Process content line by line
        lines = text_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if it's a heading
            if line.startswith('Question') or line.startswith('Plan') or \
               line.startswith('Analysis') or line.startswith('Decision') or \
               line.startswith('Conversation Log'):
                doc.add_heading(line, level=1)
            elif line.startswith('Confidence:') or line.startswith('Attempts:'):
                p = doc.add_paragraph(line)
                p.runs[0].bold = True
            else:
                doc.add_paragraph(line)
        
        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.getvalue()
    
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX export. "
            "Install it with: pip install python-docx"
        )
    except Exception as e:
        print(f"[DOCX] ❌ Conversion failed: {e}")
        return None


def create_temp_docx(html_content: str) -> Optional[str]:
    #
    # Create a temporary DOCX file from HTML content.
    #
    # Args:
    #     html_content: HTML string to convert
    #
    # Returns:
    #     Path to temporary DOCX file, or None if conversion fails
    #
    
    try:
        docx_bytes = html_to_docx(html_content)
        
        if docx_bytes is None:
            return None
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(
            mode='wb',
            suffix=".docx",
            delete=False
        )
        temp_file.write(docx_bytes)
        temp_file.close()
        
        return temp_file.name
    
    except Exception as e:
        print(f"[DOCX] ❌ Failed to create temp DOCX: {e}")
        return None

